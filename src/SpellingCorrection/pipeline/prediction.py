from datasets import Dataset
import pandas as pd


from transformers import pipeline, set_seed, AutoModelForSeq2SeqLM, AutoTokenizer
from datasets import load_dataset, load_from_disk, load_metric
import matplotlib.pyplot as plt
import pandas as pd
import torch
from tqdm import tqdm
import os


from transformers import Seq2SeqTrainer, Seq2SeqTrainingArguments


from transformers import DataCollatorForSeq2Seq
from transformers import Seq2SeqTrainingArguments
from transformers import Seq2SeqTrainer
from SpellingCorrection.entity import ModelTrainerConfig

from transformers import pipeline, set_seed
from docx import Document
from fpdf import FPDF
import PyPDF2
from pathlib import Path
import os


import os
import pandas as pd
import torch
from transformers import pipeline, AutoTokenizer, AutoModelForSeq2SeqLM
from docx import Document
from fpdf import FPDF
from pathlib import Path
import pandas as pd
from PyPDF2 import PdfReader

from SpellingCorrection.config.configuration import ConfigurationManager
class PredictionPipeline:
    def __init__(self):
        self.config = ConfigurationManager().get_model_evaluation_config()
        self.device = "cuda" if torch.cuda.is_available() else "cpu"
        self.tokenizer = AutoTokenizer.from_pretrained(self.config.tokenizer_path)
        self.model = AutoModelForSeq2SeqLM.from_pretrained(self.config.model_path).to(self.device)
        self.pipe = pipeline("text2text-generation", model=self.model, tokenizer=self.tokenizer, device=0 if torch.cuda.is_available() else -1)

    def predict(self, text):
        print("Input text: ", text)
        prediction = self.pipe(text, max_length=128, clean_up_tokenization_spaces=True)[0]['generated_text']
        print("Output text: ", prediction)
        return prediction

    def save_predictions(self, inputs, outputs, output_dir):
        if not os.path.exists(output_dir):
            os.makedirs(output_dir)

        # Save as text file
        with open(os.path.join(output_dir, 'predictions.txt'), 'w', encoding='utf-8') as f:
            for input_text, output_text in zip(inputs, outputs):
                f.write(f"Input: {input_text}\nOutput: {output_text}\n\n")

        # Save as Word file
        doc = Document()
        for input_text, output_text in zip(inputs, outputs):
            doc.add_heading('Prediction', level=1)
            doc.add_paragraph(f"Input: {input_text}")
            doc.add_paragraph(f"Output: {output_text}")
        doc.save(os.path.join(output_dir, 'predictions.docx'))

        # Save as PDF file
        pdf = FPDF()
        pdf.add_page()
        pdf.set_font("Arial", size=12)
        for input_text, output_text in zip(inputs, outputs):
            pdf.multi_cell(0, 10, f"Input: {input_text}\nOutput: {output_text}\n\n")
        pdf.output(os.path.join(output_dir, 'predictions.pdf'))

    def read_text_file(self, file_path):
        with open(file_path, 'r', encoding='utf-8') as file:
            return file.readlines()

    def read_word_file(self, file_path):
        doc = Document(file_path)
        return [para.text for para in doc.paragraphs]

    def read_pdf_file(self, file_path):
        reader = PyPDF2.PdfReader(file_path)
        text = []
        for page in reader.pages:
            text.append(page.extract_text())
        return text

    def process_file(self, file_path):
        file_extension = Path(file_path).suffix.lower()
        if file_extension == '.txt':
            return self.read_text_file(file_path)
        elif file_extension == '.docx':
            return self.read_word_file(file_path)
        elif file_extension == '.pdf':
            return self.read_pdf_file(file_path)
        else:
            raise ValueError(f"Unsupported file format: {file_extension}")

